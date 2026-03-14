import io
import fitz  # PyMuPDF
from docx import Document


class Parser:
    def __init__(self, content: bytes, content_type: str):
        self.content = content
        self.content_type = content_type

    def parse(self) -> str:
        """Parse the content based on its type and return plain text."""
        if self.content_type == "application/pdf":
            return self._parse_pdf()
        elif self.content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]:
            return self._parse_docx()
        elif self.content_type == "text/plain":
            return self._parse_txt()
        else:
            raise ValueError("Unsupported file type")

    def _parse_pdf(self) -> str:
        doc = fitz.Document(stream=self.content)
        raw_text = ""
        for page in doc:
            raw_text += page.get_text()
        return raw_text

    def _parse_docx(self) -> str:
        doc = Document(io.BytesIO(self.content))
        raw_text = ""
        for para in doc.paragraphs:
            raw_text += para.text + "\n"
        return raw_text

    def _parse_txt(self) -> str:
        return self.content.decode("utf-8")
